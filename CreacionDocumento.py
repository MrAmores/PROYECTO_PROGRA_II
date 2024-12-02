from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Adjusts the font size of all paragraphs within a cell
def ajustar_fuente(celda, tamanoFuente):
    for para in celda.paragraphs:
        para.runs[0].font.size = Pt(tamanoFuente)

def crearDocumento(objSoliServicio):
    listSS = objSoliServicio.listaWord()

    #Create the Word document
    doc = Document()
    doc.add_heading('Listado de Solicitudes de Servicio', level=1)

    if listSS:
        #Create the table
        columnas = listSS[0].keys()
        tabla = doc.add_table(rows=1, cols=len(columnas))
        tabla.style = 'Table Grid'

        #Headings
        encabezados = tabla.rows[0].cells
        for i, columna in enumerate(columnas):
            encabezados[i].text = columna
            encabezados[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ajustar_fuente(encabezados[i], 7.5)

        #Adds rows
        for fila in listSS:
            row = tabla.add_row().cells
            for i, valor in enumerate(fila.values()):
                row[i].text = str(valor)
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                ajustar_fuente(row[i], 7.5)

        # Adjust the width of columns
        anchoColumnas = [3.5, 5.0, 3.5, 2.5, 5.0, 4.0, 3.0, 3.0, 4.0, 2.5]
        for i, column in enumerate(tabla.columns):
            for cell in column.cells:
                cell.width = Pt(anchoColumnas[i] * 72)  #Convert from inches to points
    else:
        doc.add_paragraph("No hay datos disponibles.")

    #Save the document
    doc.save("ListadoSolicitudes.docx")
    print("\n=========================================================")
    print("Documento generado exitosamente: ListadoSolicitudes.docx")
    print("=========================================================\n")